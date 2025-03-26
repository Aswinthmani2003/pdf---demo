import streamlit as st
from docx import Document
from datetime import datetime
import os
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import uuid
import tempfile

# Proposal configurations for uploaded documents
PROPOSAL_CONFIG = {
    "Manychat & CRM Automation": {
        "template": "Manychat & CRM Automation.docx",
        "pricing_fields": [
            ("ManyChat Automation", "MC-Price"),
            ("CRM Automations", "C-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Make & CRM Automation": {
        "template": "Make & CRM Automation.docx",
        "pricing_fields": [
            ("Make Automation", "M-Price"),
            ("CRM Automations", "C-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Make & Manychat Automation": {
        "template": "Make & Manychat Automation.docx",
        "pricing_fields": [
            ("ManyChat Automation", "MC-Price"),
            ("Make Automation", "M-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Make, Manychat & CRM Automation": {
        "template": "Make, Manychat & CRM Automation.docx",
        "pricing_fields": [
            ("ManyChat Automation", "MC-Price"),
            ("Make Automation", "M-Price"),
            ("CRM Automations", "C-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    }
}

def apply_formatting(new_run, original_run):
    """Copy formatting from original run to new run"""
    if original_run.font.name:
        new_run.font.name = original_run.font.name
        new_run._element.rPr.rFonts.set(qn('w:eastAsia'), original_run.font.name)
    if original_run.font.size:
        new_run.font.size = original_run.font.size
    if original_run.font.color.rgb:
        new_run.font.color.rgb = original_run.font.color.rgb
    new_run.bold = original_run.bold
    new_run.italic = original_run.italic

def replace_in_paragraph(para, placeholders):
    """Handle paragraph replacements preserving formatting"""
    original_runs = para.runs.copy()
    full_text = para.text
    for ph, value in placeholders.items():
        full_text = full_text.replace(ph, str(value))

    if full_text != para.text:
        para.clear()
        new_run = para.add_run(full_text)
        if original_runs:
            original_run = next((r for r in original_runs if r.text), None)
            if original_run:
                apply_formatting(new_run, original_run)

def replace_and_format(doc, placeholders):
    """Enhanced replacement with table cell handling"""
    # Process paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para, placeholders)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.tables:
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            for nested_cell in nested_row.cells:
                                for para in nested_cell.paragraphs:
                                    replace_in_paragraph(para, placeholders)
                else:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, placeholders)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return doc

def get_general_team_details():
    """Collect team composition for non-marketing proposals"""
    st.subheader("Team Composition")
    team_roles = {
        "Project Manager": "P1",
        "Frontend Developers": "F1",
        "Business Analyst": "B1",
        "AI/ML Developers": "A1",
        "UI/UX Members": "U1",
        "System Architect": "S1",
        "Backend Developers": "BD1",
        "AWS Developer": "AD1"
    }
    team_details = {}
    cols = st.columns(2)

    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 2]:
            count = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"team_{placeholder}"
            )
            team_details[f"<<{placeholder}>>"] = str(count)
    return team_details

def remove_empty_rows(table):
    """Remove rows from the table where the second cell is empty or has no value."""
    rows_to_remove = []
    for row in table.rows:
        if len(row.cells) > 1 and row.cells[1].text.strip() == "":
            rows_to_remove.append(row)
    # Remove rows in reverse order to avoid index issues
    for row in reversed(rows_to_remove):
        table._tbl.remove(row._element)

def validate_phone_number(country, phone_number):
    """Validate phone number based on country"""
    if country.lower() == "india":
        if not phone_number.startswith("+91"):
            return False
    else:
        if not phone_number.startswith("+1"):
            return False
    return True

def format_number_with_commas(number):
    """Format number with commas (e.g., 10000 -> 10,000)"""
    return f"{number:,}"

def generate_document():
    st.title("Proposal Generator")
    base_dir = os.getcwd()

    selected_proposal = st.selectbox("Select Proposal", list(PROPOSAL_CONFIG.keys()))
    config = PROPOSAL_CONFIG[selected_proposal]
    template_path = os.path.join(base_dir, config["template"])

    # Client Information
    col1, col2 = st.columns(2)
    with col1:
        client_name = st.text_input("Client Name:")
        client_email = st.text_input("Client Email:")
    with col2:
        country = st.text_input("Country:")
        client_number = st.text_input("Client Number:")
        if client_number and country:
            if not validate_phone_number(country, client_number):
                st.error(f"Phone number for {country} should start with {'+91' if country.lower() == 'india' else '+1'}")

    date_field = st.date_input("Date:", datetime.today())

    # Currency Handling
    currency = st.selectbox("Select Currency", ["USD", "INR"])
    currency_symbol = "$" if currency == "USD" else "₹"

    # Special Fields Handling
    special_data = {}
    if config.get("special_fields"):
        st.subheader("Additional Details")
        for field, wrapper in config["special_fields"]:
            if wrapper == "<<":
                placeholder = f"<<{field}>>"
                if field == "VDate":
                    vdate = st.date_input("Proposal Validity Until:")
                    special_data[placeholder] = vdate.strftime("%d-%m-%Y")

    # Pricing Section
    st.subheader("Pricing Details")
    pricing_data = {}
    numerical_values = {}  # To store raw numerical values for calculations

    num_pricing_fields = len(config["pricing_fields"])
    cols = st.columns(num_pricing_fields)

    for idx, (label, key) in enumerate(config["pricing_fields"]):
        with cols[idx % num_pricing_fields]:
            value = st.number_input(
                f"{label} ({currency})",
                min_value=0,
                value=0,
                step=100,
                format="%d",
                key=f"price_{key}"
            )
            numerical_values[key] = value
            if value > 0:
                pricing_data[f"<<{key}>>"] = f"{currency_symbol}{format_number_with_commas(value)}"
            else:
                pricing_data[f"<<{key}>>"] = ""

    # Calculate Total Amount
    services_sum = sum(numerical_values.values())
    am_price = int(services_sum * 0.10)
    pricing_data["<<AM-Price>>"] = f"{currency_symbol}{format_number_with_commas(am_price)}"
    
    total = services_sum + am_price
    if currency == "INR":
        pricing_data["<<T-Price>>"] = f"{currency_symbol}{format_number_with_commas(total)} + 18% GST"
    else:
        pricing_data["<<T-Price>>"] = f"{currency_symbol}{format_number_with_commas(total)}"

    # Additional Features
    af_price = 250 if currency == "USD" else 25000
    pricing_data["<<AF-Price>>"] = f"{currency_symbol}{format_number_with_commas(af_price)}"

    # Team Composition
    team_data = get_general_team_details()

    # Additional Tools
    st.subheader("Add Additional Tools")
    additional_tool_1 = st.text_input("Tool 1:")
    additional_tool_2 = st.text_input("Tool 2:")

    additional_tools_data = {
        "<<T1>>": additional_tool_1 if additional_tool_1 else "",
        "<<T2>>": additional_tool_2 if additional_tool_2 else ""
    }

    # Combine all placeholders
    placeholders = {
        "<<Client Name>>": client_name,
        "<<Client Email>>": client_email,
        "<<Client Number>>": client_number,
        "<<Date>>": date_field.strftime("%d-%m-%Y"),
        "<<Country>>": country
    }
    placeholders.update(pricing_data)
    placeholders.update(team_data)
    placeholders.update(special_data)
    placeholders.update(additional_tools_data)

    if st.button("Generate Proposal"):
        if client_number and country and not validate_phone_number(country, client_number):
            st.error(f"Invalid phone number format for {country} should start with {'+91' if country.lower() == 'india' else '+1'}.")
        else:
            formatted_date = date_field.strftime("%d-%m-%Y")
            doc_filename = f"Automation Proposal - {client_name} {formatted_date}.docx"

            with tempfile.TemporaryDirectory() as temp_dir:
                try:
                    doc = Document(template_path)
                except FileNotFoundError:
                    st.error(f"Template file not found: {template_path}")
                    return

                doc = replace_and_format(doc, placeholders)

                for table in doc.tables:
                    remove_empty_rows(table)

                doc_path = os.path.join(temp_dir, doc_filename)
                doc.save(doc_path)

                with open(doc_path, "rb") as f:
                    st.download_button(
                        label="Download Proposal",
                        data=f,
                        file_name=doc_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

if __name__ == "__main__":
    generate_document()